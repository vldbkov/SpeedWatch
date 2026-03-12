# report_generator.py
from datetime import datetime
import os

class ReportGenerator:
    """Генератор детальных отчетов для провайдера"""
    
    def __init__(self, db, parent):
        self.db = db
        self.parent = parent  # ссылка на InternetSpeedMonitor для доступа к данным

    def _format_date_for_display(self, period_name, start_date, end_date):
        """Форматирование дат в европейский формат дд-мм-гггг"""
        try:
            # Если start_date и end_date уже в формате строк, пытаемся преобразовать
            if start_date != "начало" and '-' in start_date:
                # Предполагаем, что приходит в формате YYYY-MM-DD
                parts = start_date.split('-')
                if len(parts) == 3:
                    start_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
            
            if end_date != "конец" and '-' in end_date:
                parts = end_date.split('-')
                if len(parts) == 3:
                    end_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
        except:
            pass
        
        return f"📅 {period_name} · {start_date} — {end_date}"

    def generate_html_report(self, period_name, start_date, end_date, stats):
        """Генерация HTML-отчета"""
        
        # Определяем цвета для статусов
        planned = self.parent.planned_speed_var.get() if hasattr(self.parent, 'planned_speed_var') else 0
        if planned > 0:
            percent = (stats['avg_download'] / planned) * 100
            if percent >= 90:
                status_color = "green"
                status_text = "✅ Хорошо"
            elif percent >= 70:
                status_color = "orange"
                status_text = "⚠️ Средне"
            else:
                status_color = "red"
                status_text = "❌ Плохо"
        else:
            status_color = "gray"
            status_text = "ℹ️ Нет данных"
        
        # Формируем HTML
        html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SpeedWatch - Отчет качества</title>
    <style>
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .report {{
            background: white;
            border-radius: 15px;
            padding: 30px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-top: 0;
        }}
        h2 {{
            color: #34495e;
            margin-top: 25px;
            font-size: 1.3em;
        }}
        .period {{
            color: #7f8c8d;
            font-size: 0.9em;
            margin-bottom: 20px;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(2, 1fr);
            gap: 20px;
            margin: 20px 0;
        }}
        .stat-card {{
            background: #f8f9fa;
            border-radius: 10px;
            padding: 15px;
            border-left: 4px solid #3498db;
        }}
        .stat-value {{
            font-size: 1.8em;
            font-weight: bold;
            color: #2c3e50;
        }}
        .stat-label {{
            color: #7f8c8d;
            font-size: 0.9em;
            margin-top: 5px;
        }}
        .status-{status_color} {{
            color: {status_color};
            font-weight: bold;
        }}
        .warning {{
            background: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 15px;
            margin: 20px 0;
            border-radius: 5px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th {{
            background: #3498db;
            color: white;
            padding: 10px;
            text-align: left;
        }}
        td {{
            padding: 10px;
            border-bottom: 1px solid #ddd;
        }}
        tr:hover {{
            background: #f5f5f5;
        }}
        .footer {{
            margin-top: 30px;
            text-align: center;
            color: #95a5a6;
            font-size: 0.8em;
            border-top: 1px solid #ddd;
            padding-top: 20px;
        }}
    </style>
</head>
<body>
    <div class="report">
        <h1>📊 SpeedWatch - Отчет качества</h1>
        <div class="period">📅 {self._format_date_for_display(period_name, start_date, end_date)}</div>
        
        <!-- ИНФОРМАЦИЯ О ПОДКЛЮЧЕНИИ (только провайдер и IP) из данных периода -->
        <div class="connection-info" style="background: #e8f4f8; padding: 10px 15px; border-radius: 8px; margin: 15px 0; border-left: 4px solid #3498db;">
            <table style="width: 100%; border-collapse: collapse;">
                <tr>
                    <td style="width: 50%; padding: 5px;"><strong>Провайдер:</strong> {stats.get('client_provider', '—')}</td>
                    <td style="width: 50%; padding: 5px;"><strong>IP адрес:</strong> {stats.get('client_ip', '—')}</td>
                </tr>
            </table>
        </div>
        
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-value">{stats['avg_download']:.1f} Mbps</div>
                <div class="stat-label">📥 Средняя загрузка</div>
                <div>Лучшая: {stats['max_download']:.1f} · Худшая: {stats['min_download']:.1f}</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{stats['avg_upload']:.1f} Mbps</div>
                <div class="stat-label">📤 Средняя отдача</div>
            </div>
        </div>
        
        <h2>📶 Пинг и джиттер</h2>
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-value">{stats['avg_ping']:.1f} ms</div>
                <div class="stat-label">Пинг (сред.)</div>
                <div>Макс: {stats['max_ping']:.1f} · Мин: {stats['min_ping']:.1f}</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{stats['avg_jitter']:.1f} ms</div>
                <div class="stat-label">Джиттер (сред.)</div>
                <div>Макс: {stats['max_jitter']:.1f}</div>
            </div>
        </div>
"""
        
        # Добавляем информацию о тарифе
        if planned > 0:
            html += f"""
        <h2>📋 Соответствие тарифу</h2>
        <div class="stat-card">
            <div class="stat-value">{percent:.1f}%</div>
            <div class="stat-label">от заявленных {planned} Mbps</div>
            <div class="status-{status_color}">{status_text}</div>
        </div>
"""
        
        # Добавляем информацию о проблемных периодах
        if stats.get('hourly') or stats.get('daily'):
            html += """
        <h2>⏰ Проблемные периоды</h2>
        <table>
            <tr>
                <th>Период</th>
                <th>Скорость</th>
                <th>Пинг</th>
            </tr>
"""
            if stats.get('hourly'):
                # Топ-3 худших часа
                bad_hours = sorted(stats['hourly'], key=lambda x: x[1])[:3]
                for hour_data in bad_hours:
                    hour = int(hour_data[0])
                    html += f"""
            <tr>
                <td>{hour:02d}:00 — {hour+1:02d}:00</td>
                <td>{hour_data[1]:.1f} Mbps</td>
                <td>{hour_data[2]:.1f} ms</td>
            </tr>
"""
            html += "</table>"
        
        # Добавляем статистику измерений
        html += f"""
        <h2>📊 Общая статистика</h2>
        <table>
            <tr>
                <td>📊 Всего измерений</td>
                <td><strong>{stats['count']}</strong></td>
            </tr>
            <tr>
                <td>⏱️ Измерений в час (сред.)</td>
                <td><strong>{stats['measurements_per_hour']:.1f}</strong></td>
            </tr>
            <tr>
                <td>📅 Всего часов в периоде</td>
                <td><strong>{stats['total_hours']}</strong></td>
            </tr>
        </table>
        
        <div class="footer">
            📅 Отчет сгенерирован {datetime.now().strftime('%d.%m.%Y %H:%M')}<br>
            SpeedWatch — мониторинг скорости интернета
        </div>
    </div>
</body>
</html>"""
        return html
    
    def generate_docx_report(self, period_name, start_date, end_date, stats):
        """Генерация DOCX-отчета (полная версия как в HTML)"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('SpeedWatch - Отчет качества', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Период
            # Форматируем даты для DOCX
            display_start = start_date
            display_end = end_date
            try:
                if start_date != "начало" and '-' in start_date:
                    parts = start_date.split('-')
                    if len(parts) == 3:
                        display_start = f"{parts[2]}-{parts[1]}-{parts[0]}"
                if end_date != "конец" and '-' in end_date:
                    parts = end_date.split('-')
                    if len(parts) == 3:
                        display_end = f"{parts[2]}-{parts[1]}-{parts[0]}"
            except:
                pass
            
            doc.add_paragraph(f"Период: {period_name} ({display_start} — {display_end})")

            # ===== ИНФОРМАЦИЯ О ПОДКЛЮЧЕНИИ (провайдер и IP) =====
            if stats.get('client_provider') or stats.get('client_ip'):
                # Создаем таблицу с двумя колонками
                provider_table = doc.add_table(rows=1, cols=2)
                provider_table.style = 'Light Grid Accent 1'
                
                # Левая колонка - Провайдер
                left_cell = provider_table.cell(0, 0)
                left_cell.text = f"Провайдер: {stats.get('client_provider', '—')}"
                
                # Правая колонка - IP адрес
                right_cell = provider_table.cell(0, 1)
                right_cell.text = f"IP адрес: {stats.get('client_ip', '—')}"              

            # ===== СКОРОСТЬ =====
            doc.add_heading('📈 Скорость', level=1)
            
            # Таблица скорости
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Заголовки
            table.cell(0,0).text = 'Показатель'
            table.cell(0,1).text = 'Значение'
            
            # Данные
            table.cell(1,0).text = 'Средняя загрузка'
            table.cell(1,1).text = f"{stats['avg_download']:.1f} Mbps"
            
            table.cell(2,0).text = 'Лучшая загрузка'
            table.cell(2,1).text = f"{stats['max_download']:.1f} Mbps"
            
            table.cell(3,0).text = 'Худшая загрузка'
            table.cell(3,1).text = f"{stats['min_download']:.1f} Mbps"
            
            table.cell(4,0).text = 'Средняя отдача'
            table.cell(4,1).text = f"{stats['avg_upload']:.1f} Mbps"
            
            # ===== ПИНГ И ДЖИТТЕР =====
            doc.add_heading('📶 Пинг и джиттер', level=1)
            
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'
            
            table.cell(0,0).text = 'Показатель'
            table.cell(0,1).text = 'Значение'
            
            table.cell(1,0).text = 'Средний пинг'
            table.cell(1,1).text = f"{stats['avg_ping']:.1f} ms"
            
            table.cell(2,0).text = 'Максимальный пинг'
            table.cell(2,1).text = f"{stats['max_ping']:.1f} ms"
            
            table.cell(3,0).text = 'Минимальный пинг'
            table.cell(3,1).text = f"{stats['min_ping']:.1f} ms"
            
            table.cell(4,0).text = 'Средний джиттер'
            table.cell(4,1).text = f"{stats['avg_jitter']:.1f} ms"
            
            # ===== СООТВЕТСТВИЕ ТАРИФУ =====
            planned = self.parent.planned_speed_var.get() if hasattr(self.parent, 'planned_speed_var') else 0
            if planned > 0:
                doc.add_heading('📋 Соответствие тарифу', level=1)
                percent = (stats['avg_download'] / planned) * 100
                
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Light Grid Accent 1'
                
                table.cell(0,0).text = 'Заявленная скорость'
                table.cell(0,1).text = f"{planned} Mbps"
                
                table.cell(1,0).text = 'Соответствие'
                table.cell(1,1).text = f"{percent:.1f}%"
                
                if percent >= 90:
                    doc.add_paragraph("✅ Хорошо")
                elif percent >= 70:
                    doc.add_paragraph("⚠️ Средне")
                else:
                    doc.add_paragraph("❌ Плохо")
            
            # ===== ПРОБЛЕМНЫЕ ПЕРИОДЫ =====
            if stats.get('hourly') or stats.get('daily'):
                doc.add_heading('⏰ Проблемные периоды', level=1)
                
                # Создаем таблицу с двумя колонками
                problem_table = doc.add_table(rows=1, cols=2)
                problem_table.style = 'Light Grid Accent 1'
                
                # Левая колонка - Худшее время суток
                left_cell = problem_table.cell(0, 0)
                left_cell.text = "🕐 Худшее время суток"
                
                # Добавляем список в левую ячейку
                if stats.get('hourly'):
                    left_paragraph = left_cell.paragraphs[0]
                    left_paragraph.text = "🕐 Худшее время суток"
                    left_paragraph.style = 'Heading 3'
                    
                    bad_hours = sorted(stats['hourly'], key=lambda x: x[1])[:3]
                    for hour_data in bad_hours:
                        hour = int(hour_data[0])
                        p = left_cell.add_paragraph()
                        p.add_run(f"  • {hour:02d}:00 — {hour+1:02d}:00").bold = False
                        p.add_run(f"  ({hour_data[1]:.1f} Mbps)")
                else:
                    left_cell.text = "🕐 Худшее время суток\nНет данных"
                
                # Правая колонка - Худшие дни
                right_cell = problem_table.cell(0, 1)
                
                if stats.get('daily'):
                    right_cell.text = "📉 Худшие дни"
                    right_paragraph = right_cell.paragraphs[0]
                    right_paragraph.style = 'Heading 3'
                    
                    bad_days = sorted(stats['daily'], key=lambda x: x[2])[:3]
                    days = ['Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс']
                    for day_data in bad_days:
                        day_name = days[int(day_data[0])]
                        p = right_cell.add_paragraph()
                        p.add_run(f"  • {day_name}").bold = False
                        p.add_run(f" ({day_data[1][8:10]}.{day_data[1][5:7]})")
                        p.add_run(f"  {day_data[2]:.1f} Mbps")
                else:
                    right_cell.text = "📉 Худшие дни\nНет данных"
            
            # ===== ОБЩАЯ СТАТИСТИКА =====
            doc.add_heading('📊 Общая статистика', level=1)
            
            # Создаем таблицу с тремя колонками
            stats_table = doc.add_table(rows=1, cols=3)
            stats_table.style = 'Light Grid Accent 1'
            
            # Левая колонка - Всего измерений
            left_cell = stats_table.cell(0, 0)
            left_cell.text = f"📊 Всего измерений: {stats['count']}"
            
            # Средняя колонка - Измерений в час
            middle_cell = stats_table.cell(0, 1)
            if 'measurements_per_hour' in stats:
                middle_cell.text = f"⏱️ Изм/час: {stats['measurements_per_hour']:.1f}"
            else:
                middle_cell.text = "⏱️ Изм/час: —"
            
            # Правая колонка - Всего часов в периоде
            right_cell = stats_table.cell(0, 2)
            if 'total_hours' in stats:
                right_cell.text = f"📅 Всего часов: {stats['total_hours']}"
            else:
                right_cell.text = "📅 Всего часов: —"
            
            # ===== ПОДВАЛ =====
            doc.add_paragraph()
            doc.add_paragraph(f"Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph("SpeedWatch — мониторинг скорости интернета")
            
            return doc
            
        except ImportError as e:
            print(f"Ошибка импорта DOCX: {e}")
            return None
        except Exception as e:
            print(f"Ошибка создания DOCX: {e}")
            return None

    def save_html_report(self, filename, content):
        """Сохранение HTML-отчета"""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"Ошибка сохранения HTML: {e}")
            return False
    
    def save_docx_report(self, filename, doc):
        """Сохранение DOCX-отчета"""
        try:
            doc.save(filename)
            return True
        except Exception as e:
            print(f"Ошибка сохранения DOCX: {e}")
            return False